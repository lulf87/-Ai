from pathlib import Path
import hashlib
import os
import shutil
import subprocess

os.environ.setdefault("REG_REVIEW_DATABASE_PATH", "data/test.sqlite3")
os.environ.setdefault("LLM_PROVIDER", "fake")

from docx import Document
from fastapi.testclient import TestClient
import pytest
from sqlmodel import Session

from backend.app import regulations as regulation_services
from backend.app.database import engine, reset_database
from backend.app.llm import LLMProviderResult
from backend.app.main import app
from backend.app.models import RegulationRecord


def make_client():
    reset_database()
    return TestClient(app)


def create_golden_project(client: TestClient) -> dict:
    project = client.post(
        "/projects",
        json={
            "name": "微波消融黄金测试",
            "registration_scenario": "国产三类首次注册",
            "origin_type": "domestic",
            "application_type": "initial",
            "device_class": "III",
            "classification_code": "",
            "has_software": True,
            "is_networked": True,
            "has_ai": True,
            "outputs_energy": True,
            "has_disposable_accessory": True,
        },
    ).json()

    for sample in sorted(Path("samples/golden/microwave_ablation").glob("*.md")):
        response = client.post(
            f"/projects/{project['id']}/documents",
            files={"file": (sample.name, sample.read_bytes(), "text/markdown")},
            data={"document_type": sample.stem.split("_", 1)[1]},
        )
        assert response.status_code == 200
    return project


def report_text(filename: str) -> str:
    path = Path("reports/generated") / filename
    document = Document(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def test_backend_root_is_a_helpful_entrypoint():
    client = make_client()

    response = client.get("/")

    assert response.status_code == 200
    assert "注册资料预审" in response.text


def test_unverified_regulation_cannot_be_referenced_by_findings(tmp_path):
    client = make_client()

    project = client.post(
        "/projects",
        json={
            "name": "微波消融黄金测试",
            "registration_scenario": "国产三类首次注册",
            "origin_type": "domestic",
            "application_type": "initial",
            "device_class": "III",
            "classification_code": "",
            "has_software": True,
            "is_networked": True,
            "has_ai": True,
            "outputs_energy": True,
            "has_disposable_accessory": True,
        },
    ).json()

    regulation_document = Document()
    regulation_document.add_paragraph("医疗器械网络安全注册审查指导原则测试正文。")
    regulation_path = tmp_path / "network-security-guideline.docx"
    regulation_document.save(regulation_path)
    regulation = client.post(
        "/regulations/import/file",
        data={
            "title": "医疗器械网络安全注册审查指导原则",
            "reference_number": "CMDE 网络安全指导原则",
            "publication_date": "2022-03-01",
            "official_url": "https://www.cmde.org.cn/",
            "applicable_modules": "network_security",
            "coverage_classes": "II,III",
        },
        files={
            "file": (
                regulation_path.name,
                regulation_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    ).json()

    sample = Path("samples/golden/microwave_ablation/03_instructions.md")
    response = client.post(
        f"/projects/{project['id']}/documents",
        files={"file": (sample.name, sample.read_bytes(), "text/markdown")},
        data={"document_type": "instructions"},
    )
    assert response.status_code == 200

    run = client.post(f"/projects/{project['id']}/run-checks")
    assert run.status_code == 200
    findings = run.json()["findings"]
    assert findings
    assert all(f["regulation_id"] is None for f in findings)

    verified = client.patch(
        f"/regulations/{regulation['id']}/verify",
        json={"verification_status": "verified", "verified_by": "tester"},
    )
    assert verified.status_code == 200

    rerun = client.post(f"/projects/{project['id']}/run-checks")
    assert rerun.status_code == 200
    verified_findings = rerun.json()["findings"]
    assert any(f["regulation_id"] == regulation["id"] for f in verified_findings)


def test_preset_regulations_cover_class_ii_and_iii_active_modules():
    client = make_client()

    regulations = client.get("/regulations").json()

    presets = [regulation for regulation in regulations if regulation["source_type"] == "preset"]
    assert len(presets) >= 8
    assert all({"II", "III"}.issubset(set(regulation["coverage_classes"])) for regulation in presets)
    covered_modules = {
        module for regulation in presets for module in regulation["applicable_modules"]
    }
    assert {
        "general_submission",
        "testing",
        "software",
        "network_security",
        "clinical",
        "ai_algorithm",
        "reliability",
        "labeling",
    }.issubset(covered_modules)
    assert all(regulation["verification_status"] == "pending" for regulation in presets)
    assert any(regulation["source_files"] for regulation in presets)


def test_preset_official_attachments_are_seeded_as_data_sources():
    client = make_client()
    regulations = client.get("/regulations").json()
    regulation = next(
        item for item in regulations if item["title"] == "医疗器械产品技术要求编写指导原则"
    )

    response = client.get(f"/regulations/{regulation['id']}/attachments")

    assert response.status_code == 200
    attachments = response.json()
    assert len(attachments) == 1
    attachment = attachments[0]
    assert attachment["filename"] == "国家药品监督管理局2022年第8号通告附件.doc"
    assert attachment["source_type"] == "official_attachment"
    assert attachment["verification_usable"] is True
    assert attachment["sha256"] == "466c5784ecb9d5dd6675cac84bafe9f0ab3fd9e2975a7d85796d1974a4a6134f"


def test_metadata_only_preset_sha_does_not_enable_verification():
    client = make_client()
    regulations = client.get("/regulations").json()
    preset = next(
        regulation
        for regulation in regulations
        if regulation["title"] == "医疗器械注册申报资料要求和批准证明文件格式"
    )
    attachments = client.get(f"/regulations/{preset['id']}/attachments").json()
    assert len(attachments) == 9
    assert {attachment["source_type"] for attachment in attachments} == {"official_attachment"}
    assert {attachment["download_status"] for attachment in attachments} == {"metadata_only"}

    response = client.patch(
        f"/regulations/{preset['id']}/verify",
        json={"verification_status": "verified", "verified_by": "tester"},
    )

    assert response.status_code == 400
    assert "source SHA" in response.text


def test_file_import_extracts_regulation_text_and_sha(tmp_path):
    client = make_client()
    document = Document()
    document.add_heading("软件注册审查指导原则摘录", level=1)
    document.add_paragraph("本文件用于测试法规导入，包含软件版本和网络安全资料要求。")
    path = tmp_path / "software-guideline.docx"
    document.save(path)
    payload = path.read_bytes()

    response = client.post(
        "/regulations/import/file",
        data={
            "title": "测试法规文件导入",
            "official_url": "https://www.cmde.org.cn/test-guideline.html",
            "reference_number": "测试通告",
            "publication_date": "2026-05-14",
            "applicable_modules": "software,network_security",
            "coverage_classes": "II,III",
            "device_scope": "II类和III类有源医疗器械注册测试",
        },
        files={
            "file": (
                path.name,
                payload,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    regulation = response.json()
    expected_sha = hashlib.sha256(payload).hexdigest()
    assert regulation["attachment_sha256"] == expected_sha
    assert regulation["source_type"] == "file_import"
    assert regulation["coverage_classes"] == ["II", "III"]
    assert regulation["segment_count"] == 1
    assert "软件版本" in regulation["text_preview"]

    attachments = client.get(f"/regulations/{regulation['id']}/attachments").json()
    assert len(attachments) == 1
    attachment = attachments[0]
    assert attachment["filename"] == path.name
    assert attachment["sha256"] == expected_sha
    assert attachment["source_type"] == "uploaded_file"
    assert attachment["verification_usable"] is True
    assert attachment["segment_count"] == 1

    segments = client.get(f"/regulations/{regulation['id']}/segments").json()
    assert segments[0]["locator"] == "软件注册审查指导原则摘录"
    assert segments[0]["attachment_id"] == attachment["id"]
    assert "网络安全资料要求" in segments[0]["text"]

    search = client.get("/regulations/search", params={"query": "网络安全资料要求"}).json()
    assert any(result["regulation_id"] == regulation["id"] for result in search)
    assert any(result["attachment_sha256"] == expected_sha for result in search)

    verified = client.patch(
        f"/regulations/{regulation['id']}/verify",
        json={"verification_status": "verified", "verified_by": "tester"},
    )
    assert verified.status_code == 200


@pytest.mark.skipif(shutil.which("textutil") is None, reason="textutil is required to create a legacy .doc fixture")
def test_legacy_doc_regulation_attachment_is_extracted_and_searchable(tmp_path):
    client = make_client()
    source = tmp_path / "legacy-source.txt"
    source.write_text("医疗器械软件版本和网络安全资料要求\n", encoding="utf-8")
    legacy_doc = tmp_path / "legacy-guideline.doc"
    subprocess.run(
        ["textutil", "-convert", "doc", "-output", str(legacy_doc), str(source)],
        check=True,
        capture_output=True,
        text=True,
    )
    payload = legacy_doc.read_bytes()

    response = client.post(
        "/regulations/import/file",
        data={
            "title": "测试旧版 DOC 法规导入",
            "official_url": "https://www.nmpa.gov.cn/test-legacy-doc.html",
            "reference_number": "测试旧版 DOC",
            "publication_date": "2026-05-14",
            "applicable_modules": "software,network_security",
            "coverage_classes": "II,III",
            "device_scope": "II类和III类有源医疗器械注册测试",
        },
        files={"file": (legacy_doc.name, payload, "application/msword")},
    )

    assert response.status_code == 200, response.text
    regulation = response.json()
    expected_sha = hashlib.sha256(payload).hexdigest()
    assert regulation["attachment_sha256"] == expected_sha
    assert regulation["segment_count"] == 1

    search = client.get("/regulations/search", params={"query": "网络安全资料要求"}).json()
    assert any(result["regulation_id"] == regulation["id"] for result in search)


@pytest.mark.skipif(shutil.which("textutil") is None, reason="textutil is required to create a legacy .doc fixture")
def test_metadata_attachment_can_be_downloaded_and_extracted(monkeypatch, tmp_path):
    client = make_client()
    source = tmp_path / "metadata-source.txt"
    source.write_text("产品技术要求应包含检验方法和性能指标。\n", encoding="utf-8")
    legacy_doc = tmp_path / "official-attachment.doc"
    subprocess.run(
        ["textutil", "-convert", "doc", "-output", str(legacy_doc), str(source)],
        check=True,
        capture_output=True,
        text=True,
    )
    payload = legacy_doc.read_bytes()
    expected_sha = hashlib.sha256(payload).hexdigest()
    regulation = client.post(
        "/regulations",
        json={
            "title": "测试元数据附件下载",
            "reference_number": "测试元数据附件",
            "publication_date": "2026-05-14",
            "official_url": "https://www.nmpa.gov.cn/test-page.html",
            "applicable_modules": ["testing"],
            "coverage_classes": ["II", "III"],
        },
    ).json()
    with Session(engine) as session:
        record = session.get(RegulationRecord, regulation["id"])
        assert record is not None
        attachment = regulation_services.create_attachment(
            session,
            record,
            filename=legacy_doc.name,
            source_url="https://www.nmpa.gov.cn/test-attachment.doc",
            source_page_url=record.official_url,
            source_type="official_attachment",
            verification_usable=True,
            sha256=expected_sha,
            download_status="metadata_only",
        )
        session.commit()
        attachment_id = attachment.id

    monkeypatch.setattr(
        regulation_services,
        "download_attachment",
        lambda url, filename="", referer="": regulation_services.DownloadedAttachment(
            filename=filename or legacy_doc.name,
            content=payload,
            content_type="application/msword",
        ),
    )

    response = client.post(f"/regulations/{regulation['id']}/attachments/{attachment_id}/download")

    assert response.status_code == 200, response.text
    attachment = response.json()
    assert attachment["sha256"] == expected_sha
    assert attachment["download_status"] == "extracted"
    assert attachment["segment_count"] == 1
    search = client.get("/regulations/search", params={"query": "检验方法"}).json()
    assert any(result["attachment_id"] == attachment_id for result in search)


@pytest.mark.skipif(shutil.which("textutil") is None, reason="textutil is required to create a legacy .doc fixture")
def test_preset_metadata_attachments_can_be_bulk_downloaded(monkeypatch, tmp_path):
    client = make_client()
    source = tmp_path / "bulk-source.txt"
    source.write_text("预置法规附件正文用于批量下载和抽取。\n", encoding="utf-8")
    legacy_doc = tmp_path / "bulk-official-attachment.doc"
    subprocess.run(
        ["textutil", "-convert", "doc", "-output", str(legacy_doc), str(source)],
        check=True,
        capture_output=True,
        text=True,
    )
    payload = legacy_doc.read_bytes()
    expected_sha = hashlib.sha256(payload).hexdigest()
    regulation = client.post(
        "/regulations",
        json={
            "title": "测试预置附件批量下载",
            "reference_number": "测试批量下载",
            "publication_date": "2026-05-14",
            "official_url": "https://www.nmpa.gov.cn/test-bulk-page.html",
            "source_type": "preset",
            "applicable_modules": ["testing"],
            "coverage_classes": ["II", "III"],
        },
    ).json()
    with Session(engine) as session:
        record = session.get(RegulationRecord, regulation["id"])
        assert record is not None
        attachment = regulation_services.create_attachment(
            session,
            record,
            filename=legacy_doc.name,
            source_url="https://www.nmpa.gov.cn/test-bulk-attachment.doc",
            source_page_url=record.official_url,
            source_type="official_attachment",
            verification_usable=True,
            sha256=expected_sha,
            download_status="metadata_only",
        )
        session.commit()
        attachment_id = attachment.id

    monkeypatch.setattr(
        regulation_services,
        "download_attachment",
        lambda url, filename="", referer="": regulation_services.DownloadedAttachment(
            filename=filename or legacy_doc.name,
            content=payload,
            content_type="application/msword",
        ),
    )

    response = client.post("/regulations/preset-attachments/download")

    assert response.status_code == 200, response.text
    result = response.json()
    assert result["downloaded"] >= 1
    assert any(item["attachment_id"] == attachment_id and item["status"] == "downloaded" for item in result["results"])
    attachments = client.get(f"/regulations/{regulation['id']}/attachments").json()
    attachment = next(item for item in attachments if item["id"] == attachment_id)
    assert attachment["download_status"] == "extracted"
    assert attachment["segment_count"] == 1
    search = client.get("/regulations/search", params={"query": "批量下载"}).json()
    assert any(item["attachment_id"] == attachment_id for item in search)


def test_preset_without_attachments_downloads_official_web_page(monkeypatch):
    client = make_client()
    regulation = client.post(
        "/regulations",
        json={
            "title": "测试无附件网页正文预置法规",
            "reference_number": "测试网页正文",
            "publication_date": "2026-05-14",
            "official_url": "https://www.nmpa.gov.cn/test-web-page.html",
            "source_type": "preset",
            "applicable_modules": ["labeling"],
            "coverage_classes": ["II", "III"],
        },
    ).json()

    def fake_fetch_web_regulation(url: str) -> regulation_services.WebRegulationSource:
        assert url == "https://www.nmpa.gov.cn/test-web-page.html"
        return regulation_services.WebRegulationSource(
            title="测试无附件网页正文预置法规",
            text="官方网页正文用于说明书标签管理要求。",
            content_sha256="webpagesha256",
            content="<html>官方网页正文用于说明书标签管理要求。</html>".encode(),
            content_type="text/html; charset=utf-8",
        )

    monkeypatch.setattr(regulation_services, "fetch_web_regulation", fake_fetch_web_regulation)

    response = client.post("/regulations/preset-attachments/download")

    assert response.status_code == 200, response.text
    result = response.json()
    assert result["downloaded"] >= 1
    assert any(
        item["regulation_id"] == regulation["id"]
        and item["source_type"] == "web_page"
        and item["status"] == "downloaded"
        for item in result["results"]
    )
    attachments = client.get(f"/regulations/{regulation['id']}/attachments").json()
    assert len(attachments) == 1
    attachment = attachments[0]
    assert attachment["source_type"] == "web_page"
    assert attachment["verification_usable"] is True
    assert attachment["sha256"] == "webpagesha256"
    assert attachment["download_status"] == "extracted"
    assert attachment["segment_count"] == 1
    search = client.get("/regulations/search", params={"query": "说明书标签管理要求"}).json()
    assert any(item["attachment_id"] == attachment["id"] for item in search)


def test_web_import_stays_pending_without_file_sha(monkeypatch):
    client = make_client()

    def fake_fetch_web_regulation(url: str) -> regulation_services.WebRegulationSource:
        assert url == "https://example.test/regulation"
        return regulation_services.WebRegulationSource(
            title="网页法规标题",
            text="网页法规正文，包含有源医疗器械软件资料要求。",
            content_sha256="abc123",
        )

    monkeypatch.setattr(regulation_services, "fetch_web_regulation", fake_fetch_web_regulation)

    response = client.post(
        "/regulations/import/web",
        json={
            "url": "https://example.test/regulation",
            "applicable_modules": ["software"],
            "coverage_classes": ["II", "III"],
        },
    )

    assert response.status_code == 200
    regulation = response.json()
    assert regulation["title"] == "网页法规标题"
    assert regulation["source_type"] == "web_import"
    assert regulation["source_content_sha256"] == "abc123"
    assert regulation["segment_count"] == 1

    verified = client.patch(
        f"/regulations/{regulation['id']}/verify",
        json={"verification_status": "verified", "verified_by": "tester"},
    )
    assert verified.status_code == 400
    assert "source SHA" in verified.text


def test_master_data_endpoint_returns_empty_record_before_extraction():
    client = make_client()
    project = client.post(
        "/projects",
        json={
            "name": "空主数据项目",
            "registration_scenario": "国产三类首次注册",
        },
    ).json()

    response = client.get(f"/projects/{project['id']}/master-data")

    assert response.status_code == 200
    body = response.json()
    assert body["project_id"] == project["id"]
    assert body["product_name"] == ""


def test_project_can_load_golden_sample_documents_from_ui_flow():
    client = make_client()
    project = client.post(
        "/projects",
        json={
            "name": "样例加载项目",
            "registration_scenario": "国产三类首次注册",
        },
    ).json()

    response = client.post(f"/projects/{project['id']}/sample-documents/golden-microwave")

    assert response.status_code == 200
    documents = response.json()
    assert len(documents) >= 7
    assert {doc["parse_status"] for doc in documents} == {"parsed"}


def test_golden_sample_load_is_idempotent():
    client = make_client()
    project = client.post(
        "/projects",
        json={
            "name": "样例重复加载项目",
            "registration_scenario": "国产三类首次注册",
        },
    ).json()

    first = client.post(f"/projects/{project['id']}/sample-documents/golden-microwave")
    second = client.post(f"/projects/{project['id']}/sample-documents/golden-microwave")
    listed = client.get(f"/projects/{project['id']}/documents")

    assert first.status_code == 200
    assert second.status_code == 200
    assert listed.status_code == 200
    first_documents = first.json()
    second_documents = second.json()
    listed_documents = listed.json()
    assert len(first_documents) == len(second_documents) == len(listed_documents)
    assert {doc["id"] for doc in second_documents} == {doc["id"] for doc in first_documents}
    assert len(
        {
            (doc["document_type"], doc["filename"], doc["sha256"])
            for doc in listed_documents
        }
    ) == len(listed_documents)


def test_golden_dataset_identifies_at_least_fourteen_seeded_issues(tmp_path):
    client = make_client()
    project = client.post(
        "/projects",
        json={
            "name": "微波消融黄金测试",
            "registration_scenario": "国产三类首次注册",
            "origin_type": "domestic",
            "application_type": "initial",
            "device_class": "III",
            "classification_code": "",
            "has_software": True,
            "is_networked": True,
            "has_ai": True,
            "outputs_energy": True,
            "has_disposable_accessory": True,
        },
    ).json()

    for sample in sorted(Path("samples/golden/microwave_ablation").glob("*.md")):
        client.post(
            f"/projects/{project['id']}/documents",
            files={"file": (sample.name, sample.read_bytes(), "text/markdown")},
            data={"document_type": sample.stem.split("_", 1)[1]},
        )

    extraction = client.post(f"/projects/{project['id']}/extract-master-data")
    assert extraction.status_code == 200
    master = extraction.json()
    assert master["product_name"] == "微波消融治疗系统"
    assert master["software_version"] == "V1.0.3"

    run = client.post(f"/projects/{project['id']}/run-checks")
    assert run.status_code == 200
    findings = run.json()["findings"]
    must_find = {
        "R-NAME-CONSISTENCY",
        "R-MODEL-CONSISTENCY",
        "R-INTENDED-USE",
        "R-SOFTWARE-VERSION",
        "R-NETWORK-SECURITY",
        "R-TEST-COVERAGE",
        "R-STANDARD-VERSION",
        "R-SERVICE-LIFE",
        "R-DISPOSABLE-CONFLICT",
        "R-CONTRAINDICATION",
        "R-STRUCTURE-CONSISTENCY",
        "R-UNSUPPORTED-CLAIM",
        "R-REPRESENTATIVE-MODEL",
        "R-AI-ALGORITHM",
    }
    found_rule_ids = {finding["rule_id"] for finding in findings}
    assert must_find.issubset(found_rule_ids)
    assert len(findings) >= 14
    name_finding = next(
        finding for finding in findings if finding["rule_id"] == "R-NAME-CONSISTENCY"
    )
    assert "综述资料（01_overview.md" in name_finding["evidence_quote"]
    assert "产品名称：微波消融治疗系统" in name_finding["evidence_quote"]
    assert "使用说明书（03_instructions.md" in name_finding["evidence_quote"]
    assert "产品名称：肝脏微波消融系统" in name_finding["evidence_quote"]

    report = client.post(f"/projects/{project['id']}/reports")
    assert report.status_code == 200
    body = report.json()
    assert body["filename"].endswith(".docx")
    assert "能否获批" not in body["summary"]


def test_ai_extract_records_desensitized_llm_run_and_field_evidence():
    client = make_client()
    project = create_golden_project(client)

    response = client.post(f"/projects/{project['id']}/ai/extract-master-data")

    assert response.status_code == 200
    body = response.json()
    assert body["master_data"]["product_name"] == "微波消融治疗系统"
    assert body["master_data"]["software_version"] == "V1.0.3"
    assert any(span["field_name"] == "product_name" for span in body["evidence_spans"])
    assert body["llm_run"]["task_type"] == "extract_master_data"
    assert body["llm_run"]["contains_sensitive_content"] is False
    assert "植入式心脏起搏器患者禁用" not in body["llm_run"]["input_summary"]


def test_ai_extract_normalizes_structured_field_values(monkeypatch):
    class StructuredFieldProvider:
        provider_name = "structured-test"
        model_name = "structured-test-model"

        def extract_master_data(self, segments, fallback_master_data):
            return LLMProviderResult(
                output_json={
                    "fields": {
                        "model_specifications": ["A100", "A200"],
                        "structure_composition": [
                            "主机",
                            "微波消融探针",
                            "脚踏开关",
                        ],
                        "applicable_standards": ["GB 9706.1-2007", "YY 0505-2012"],
                        "is_networked": "是",
                        "outputs_energy": "输出",
                        "has_disposable_accessory": ["否"],
                    },
                    "evidence": [],
                },
                output_text="返回数组字段用于回归测试。",
                provider=self.provider_name,
                model_name=self.model_name,
                model_config={},
            )

    client = make_client()
    project = create_golden_project(client)
    monkeypatch.setattr(
        "backend.app.ai_services.get_llm_provider",
        lambda: StructuredFieldProvider(),
    )

    response = client.post(f"/projects/{project['id']}/ai/extract-master-data")

    assert response.status_code == 200, response.text
    master = response.json()["master_data"]
    assert master["model_specifications"] == "A100、A200"
    assert master["structure_composition"] == "主机、微波消融探针、脚踏开关"
    assert master["applicable_standards"] == "GB 9706.1-2007、YY 0505-2012"
    assert master["is_networked"] is True
    assert master["outputs_energy"] is True
    assert master["has_disposable_accessory"] is False


def test_ai_candidates_require_review_and_report_filters_pending_and_rejected_items():
    client = make_client()
    project = create_golden_project(client)
    client.post(f"/projects/{project['id']}/extract-master-data")
    client.post(f"/projects/{project['id']}/run-checks")

    ai_response = client.post(f"/projects/{project['id']}/ai/analyze-risks")

    assert ai_response.status_code == 200
    candidates = ai_response.json()["findings"]
    assert len(candidates) >= 2
    assert {finding["source_type"] for finding in candidates} == {"llm_candidate"}
    assert {finding["review_status"] for finding in candidates} == {"pending_review"}

    first_report = client.post(f"/projects/{project['id']}/reports").json()
    first_text = report_text(first_report["filename"])
    assert "智能辅助分析，非最终注册结论，需人工复核" in first_text
    assert candidates[0]["title"] not in first_text

    confirmed = client.post(
        f"/findings/{candidates[0]['id']}/review",
        json={"review_status": "confirmed"},
    )
    rejected = client.post(
        f"/findings/{candidates[1]['id']}/review",
        json={"review_status": "rejected"},
    )
    assert confirmed.status_code == 200
    assert rejected.status_code == 200

    final_report = client.post(f"/projects/{project['id']}/reports").json()
    final_text = report_text(final_report["filename"])
    assert candidates[0]["title"] in final_text
    assert candidates[1]["title"] not in final_text


def import_regulation_docx(
    client: TestClient,
    tmp_path: Path,
    *,
    filename: str,
    title: str,
    body: str,
    modules: str,
    verify: bool,
) -> dict:
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph(body)
    path = tmp_path / filename
    document.save(path)
    response = client.post(
        "/regulations/import/file",
        data={
            "title": title,
            "official_url": f"https://www.nmpa.gov.cn/{filename}.html",
            "reference_number": "测试法规",
            "publication_date": "2026-05-15",
            "applicable_modules": modules,
            "coverage_classes": "II,III",
            "device_scope": "II类和III类有源医疗器械注册测试",
        },
        files={
            "file": (
                path.name,
                path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 200, response.text
    regulation = response.json()
    if verify:
        verified = client.patch(
            f"/regulations/{regulation['id']}/verify",
            json={"verification_status": "verified", "verified_by": "tester"},
        )
        assert verified.status_code == 200, verified.text
        regulation = verified.json()
    return regulation


def test_regulatory_rag_uses_only_verified_sha_backed_sources(tmp_path):
    client = make_client()
    project = create_golden_project(client)
    client.post(f"/projects/{project['id']}/extract-master-data")
    unverified = import_regulation_docx(
        client,
        tmp_path,
        filename="unverified-ai-guideline.docx",
        title="未校验人工智能法规",
        body="未校验来源写明自动推荐消融参数必须立即补充算法资料，但本测试不得采用该来源。",
        modules="ai_algorithm",
        verify=False,
    )
    verified = import_regulation_docx(
        client,
        tmp_path,
        filename="verified-ai-guideline.docx",
        title="已校验人工智能医疗器械指导原则",
        body="人工智能医疗器械的软件算法研究资料应包括算法基本信息、训练验证数据、性能评价和风险控制。自动推荐消融参数需明确算法资料边界。",
        modules="ai_algorithm",
        verify=True,
    )

    response = client.post(f"/projects/{project['id']}/regulatory-rag-review")

    assert response.status_code == 200, response.text
    body = response.json()
    findings = body["findings"]
    assert findings
    assert body["llm_run"]["task_type"] == "regulatory_rag_review"
    assert {finding["source_type"] for finding in findings} == {"regulatory_rag_candidate"}
    assert {finding["review_status"] for finding in findings} == {"pending_review"}
    assert all(finding["regulation_id"] == verified["id"] for finding in findings)
    assert all(finding["regulation_id"] != unverified["id"] for finding in findings)
    assert all(finding["regulation_attachment_sha256"] == verified["attachment_sha256"] for finding in findings)


def test_regulatory_rag_candidates_include_regulation_evidence_sha(tmp_path):
    client = make_client()
    project = create_golden_project(client)
    regulation = import_regulation_docx(
        client,
        tmp_path,
        filename="clinical-guideline.docx",
        title="已校验临床评价指导原则",
        body="临床评价资料应说明同品种对比、临床证据充分性和免于临床评价路径依据，不能仅凭少量单臂数据作出充分结论。",
        modules="clinical",
        verify=True,
    )

    response = client.post(f"/projects/{project['id']}/regulatory-rag-review")

    assert response.status_code == 200, response.text
    finding = response.json()["findings"][0]
    assert finding["regulation_id"] == regulation["id"]
    assert finding["regulation_title"] == regulation["title"]
    assert finding["regulation_attachment_id"] is not None
    assert finding["regulation_attachment_filename"] == "clinical-guideline.docx"
    assert finding["regulation_attachment_sha256"] == regulation["attachment_sha256"]
    assert finding["regulation_evidence_locator"]
    assert "临床评价资料" in finding["regulation_evidence_quote"]


def test_regulatory_rag_returns_empty_without_verified_sources(tmp_path):
    client = make_client()
    project = create_golden_project(client)
    import_regulation_docx(
        client,
        tmp_path,
        filename="pending-labeling-guideline.docx",
        title="待校验说明书指导原则",
        body="说明书和标签应明确一次性使用附件处置要求。",
        modules="labeling",
        verify=False,
    )

    response = client.post(f"/projects/{project['id']}/regulatory-rag-review")

    assert response.status_code == 200, response.text
    body = response.json()
    assert body["findings"] == []
    assert body["llm_run"]["task_type"] == "regulatory_rag_review"
    assert "无可用已校验法规来源" in body["llm_run"]["output_json"]["notes"]


def test_regulation_ai_impact_stays_draft_until_manual_verification():
    client = make_client()
    regulation = client.post(
        "/regulations",
        json={
            "title": "有源医疗器械软件注册审查指导原则",
            "reference_number": "CMDE 软件指导原则",
            "publication_date": "2025-01-01",
            "official_url": "https://www.cmde.org.cn/",
            "attachment_filename": "software-guideline.pdf",
            "attachment_sha256": "def456",
            "applicable_modules": ["software"],
        },
    ).json()

    response = client.post(f"/regulations/{regulation['id']}/ai/summarize-impact")

    assert response.status_code == 200
    draft = response.json()
    assert draft["verification_status"] == "pending_review"
    assert "software" in draft["impacted_modules"]
    regulations = client.get("/regulations").json()
    assert regulations[0]["verification_status"] == "pending"
