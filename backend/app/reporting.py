from __future__ import annotations

from collections import Counter
from pathlib import Path

from docx import Document
from sqlmodel import Session, select

from backend.app.config import REPORT_DIR
from backend.app.dashboard import build_dashboard
from backend.app.models import Finding, ProductMasterData, Project, Report, utc_now


def display_text(value: str) -> str:
    return value.replace("AI 功能", "智能算法功能").replace("AI", "智能算法")


def create_report(session: Session, project_id: int) -> Report:
    project = session.get(Project, project_id)
    if project is None:
        raise ValueError("Project not found")

    master = session.exec(
        select(ProductMasterData).where(ProductMasterData.project_id == project_id)
    ).first()
    findings = reportable_findings(
        session.exec(select(Finding).where(Finding.project_id == project_id)).all()
    )
    dashboard = build_dashboard(session, project_id)
    summary = dashboard["boss_summary"]

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"project-{project_id}-registration-risk-report.docx"
    path = REPORT_DIR / filename
    write_docx(path, project, master, findings, summary, dashboard)

    report = Report(
        project_id=project_id,
        filename=filename,
        stored_path=str(path),
        summary=summary,
        created_at=utc_now(),
    )
    session.add(report)
    session.commit()
    session.refresh(report)
    return report


def reportable_findings(findings: list[Finding]) -> list[Finding]:
    return [
        finding
        for finding in findings
        if finding.source_type in {"rule", "rule_llm_confirmed", "manual"}
        or finding.review_status in {"confirmed", "edited"}
    ]


def build_summary(findings: list[Finding]) -> str:
    counts = Counter(finding.risk_level for finding in findings)
    return (
        f"本次资料预审发现 {counts.get('red', 0)} 项高风险、"
        f"{counts.get('yellow', 0)} 项中风险、{counts.get('green', 0)} 项低风险。"
        "请优先处理红色风险，并对待确认事项进行人工复核。"
    )


def write_docx(
    path: Path,
    project: Project,
    master: ProductMasterData | None,
    findings: list[Finding],
    summary: str,
    dashboard: dict,
) -> None:
    document = Document()
    document.add_heading("三类有源医疗器械注册资料预审报告", level=0)

    document.add_heading("一、老板摘要", level=1)
    document.add_paragraph("智能辅助分析，非最终注册结论，需人工复核。")
    add_boss_summary(document, dashboard, summary)

    document.add_heading("二、产品主数据表", level=1)
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "字段"
    table.rows[0].cells[1].text = "内容"
    rows = {
        "项目名称": project.name,
        "注册场景": project.registration_scenario,
        "产品名称": master.product_name if master else "",
        "型号规格": master.model_specifications if master else "",
        "适用范围": master.intended_use if master else "",
        "软件版本": master.software_version if master else "",
        "联网情况": "是" if (master and master.is_networked) else "否",
        "能量类型": master.energy_type if master else "",
        "检验型号": master.tested_model if master else "",
        "适用标准": master.applicable_standards if master else "",
    }
    for key, value in rows.items():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value

    sections = [
        ("三、资料完整性检查", ["R-NETWORK-SECURITY", "R-AI-ALGORITHM", "R-SERVICE-LIFE"]),
        ("四、资料一致性检查", ["R-NAME-CONSISTENCY", "R-MODEL-CONSISTENCY", "R-STRUCTURE-CONSISTENCY", "R-INTENDED-USE", "R-SOFTWARE-VERSION"]),
        ("五、检测与产品技术要求风险", ["R-TEST-COVERAGE", "R-REPRESENTATIVE-MODEL", "R-STANDARD-VERSION"]),
        ("六、软件/网络安全/智能算法风险", ["R-SOFTWARE-VERSION", "R-NETWORK-SECURITY", "R-AI-ALGORITHM"]),
        ("七、临床路径风险", ["R-CLINICAL-OVERCLAIM", "R-UNSUPPORTED-CLAIM"]),
    ]
    for title, rule_ids in sections:
        document.add_heading(title, level=1)
        matching = [finding for finding in findings if finding.rule_id in rule_ids]
        add_findings(document, matching)

    document.add_heading("八、红黄绿问题清单", level=1)
    add_findings(document, findings)

    document.add_heading("九、待确认事项", level=1)
    pending = [
        finding
        for finding in findings
        if finding.confidence_status != "evidence_based"
        or finding.review_status in {"pending_review", "edited"}
    ]
    if pending:
        add_findings(document, pending)
    else:
        document.add_paragraph("暂无系统标记的待确认事项。")

    document.add_heading("十、下一步行动建议", level=1)
    for finding in findings[:8]:
        document.add_paragraph(f"{finding.title}：{finding.recommended_action}", style="List Bullet")

    document.save(path)


def add_boss_summary(document: Document, dashboard: dict, summary: str) -> None:
    risk_counts = dashboard.get("risk_counts", {})
    red = risk_counts.get("red", 0)
    yellow = risk_counts.get("yellow", 0)
    overall = "红" if red else "黄" if yellow else "绿"
    document.add_paragraph(summary)
    document.add_paragraph(f"项目总体风险：{overall}")
    document.add_paragraph(f"申报准备度：{dashboard.get('readiness_score', 0)}/100")
    document.add_paragraph(
        f"重大申报断点：{red} 项；预计发补高风险项：{red} 项；需关注项：{yellow} 项。"
    )

    document.add_paragraph("责任人分布")
    owner_counts = dashboard.get("owner_counts", {})
    if owner_counts:
        for owner, count in owner_counts.items():
            document.add_paragraph(f"{owner}：{count} 项", style="List Bullet")
    else:
        document.add_paragraph("暂无责任人分布。")

    document.add_paragraph("重大申报断点")
    major_breakpoints = dashboard.get("major_breakpoints", [])
    if major_breakpoints:
        for finding in major_breakpoints[:5]:
            document.add_paragraph(
                f"{display_text(finding.title)}：{display_text(finding.recommended_action)}",
                style="List Bullet",
            )
    else:
        document.add_paragraph("暂无红色断点。")

    document.add_paragraph("下一步必须完成的动作")
    next_actions = dashboard.get("next_actions", [])
    if next_actions:
        for action in next_actions:
            document.add_paragraph(
                f"{action['owner']}｜{display_text(action['title'])}：{display_text(action['action'])}",
                style="List Bullet",
            )
    else:
        document.add_paragraph("暂无下一步动作。")


def add_findings(document: Document, findings: list[Finding]) -> None:
    if not findings:
        document.add_paragraph("未发现该类风险。")
        return
    for finding in findings:
        source_label = {
            "rule": "规则发现",
            "llm_candidate": "智能候选",
            "regulatory_rag_candidate": "法规RAG候选",
            "rule_llm_confirmed": "规则与智能辅助确认",
            "manual": "人工录入",
        }.get(finding.source_type, finding.source_type)
        review_label = {
            "pending_review": "待人工确认",
            "confirmed": "已确认",
            "rejected": "已驳回",
            "edited": "已修改确认",
        }.get(finding.review_status, finding.review_status)
        risk_label = {"red": "高风险", "yellow": "需关注", "green": "通过"}.get(
            finding.risk_level, finding.risk_level
        )
        document.add_paragraph(
            f"[{risk_label}] [{source_label}/{review_label}] {display_text(finding.title)}",
            style="List Bullet",
        )
        document.add_paragraph(f"问题：{display_text(finding.description)}")
        evidence = display_text(finding.evidence_quote) or "暂无可展示证据，请人工补充资料后复核。"
        document.add_paragraph("资料依据：")
        for line in evidence.splitlines():
            if line.strip():
                document.add_paragraph(line.strip(), style="List Bullet")
        if finding.regulation_evidence_quote:
            regulation_label = finding.regulation_title or "已校验法规"
            sha_label = (
                f"；附件 SHA：{finding.regulation_attachment_sha256}"
                if finding.regulation_attachment_sha256
                else ""
            )
            locator_label = (
                f"；位置：{finding.regulation_evidence_locator}"
                if finding.regulation_evidence_locator
                else ""
            )
            document.add_paragraph(f"法规依据：{display_text(regulation_label)}{sha_label}{locator_label}")
            document.add_paragraph(display_text(finding.regulation_evidence_quote), style="List Bullet")
        if finding.ai_rationale:
            document.add_paragraph(f"智能分析理由：{display_text(finding.ai_rationale)}")
        document.add_paragraph(f"建议：{display_text(finding.recommended_action)}")
